import docx

doc_path = "/Users/noorullah/Desktop/FYP/output/doc/FYP DOC NOORULLAH - FINAL FORMATTED - PHYSICAL ERD UPDATED.docx"
doc = docx.Document(doc_path)

full_text = "\n".join([p.text for p in doc.paragraphs])
full_text_lower = full_text.lower()

stale_terms = ["uuid", "integer", "int", "foreign key"]
for term in stale_terms:
    print(f"'{term}' appears {full_text_lower.count(term)} times in the text.")

table_text = []
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            table_text.append(cell.text)
full_table_text_lower = "\n".join(table_text).lower()

for term in stale_terms:
    print(f"'{term}' appears {full_table_text_lower.count(term)} times in the tables.")

import docx

doc_path = "/Users/noorullah/Desktop/FYP/output/doc/FYP DOC NOORULLAH - FINAL FORMATTED - PHYSICAL ERD UPDATED.docx"
doc = docx.Document(doc_path)

full_text = "\n".join([p.text for p in doc.paragraphs])
full_text_lower = full_text.lower()

if "varchar" in full_text_lower:
    print("'varchar' is present in the document text.")
else:
    print("'varchar' is MISSING from the document text.")
    
# Let's also check for table contents in the docx, as VARCHAR might be in a data dictionary table
table_text = []
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            table_text.append(cell.text)

full_table_text_lower = "\n".join(table_text).lower()
if "varchar" in full_table_text_lower:
    print("'varchar' is present in the document tables.")
else:
    print("'varchar' is MISSING from the document tables.")
